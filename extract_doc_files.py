# from docx import Document
# from datetime import datetime
# from pathlib import Path 
# import re 
# import json 



# def process_single_docx(input_file: Path, output_file: Path, verbose=False):
#     #  process a single docx file
#     try:
#         if verbose:
#             print(f"Loading docx file: {input_file}")
        
#         doc = Document(input_file)


#     except Exception as e:
#         print(f"Error loading docx file {input_file}: {e}")

# def show_docx_props(doc: Document):

#     props = doc.core_properties

#     print(props.author)
#     print(props.title)
#     print(props.created)
#     print(props.last_modified_by)
#     print(props.subject)
#     print(props.keywords)

# files = [file for file in Path(".").glob("*.docx")]
# for file in files:
#     doc = Document(file)
    # print(file)
    # print(f"Properties for file: {file}")
    # show_docx_props(doc)
    # print("-" * 40)

    # for section in doc.sections:

# for para in doc.paragraphs:
#     print(para.text, para.style.name)
#     for run in para.runs:
#         data.append({
#             "text": run.text,
#             "bold": run.bold,
#             "italic": run.italic,
#             "under_line": run.underline
#         })

# print(data)


# import mammoth
# from pathlib import Path

# files = [file for file in Path(".").glob("*.docx")]
# for file in files:
#     print(file)

#     with open(file, "rb") as docx_file:
#         result = mammoth.convert_to_html(docx_file)
#         html = result.value
    

#     filepath = f"{file.stem}.html"
#     with open(filepath, "w", encoding="utf-8") as html_file:
#         html_file.write(html)


from docx import Document
import re
from pathlib import Path


def paragraph_with_styles(doc: Document):
    out = []
    for i, para in enumerate(doc.paragraphs):
        style = None

        style = para.style.name

        out.append({
            "index": i,
            "text": para.text,
            "style": style
        })

    return out

files = [file for file in Path(".").glob("*.docx")]
for file in files:
    doc = Document(file)
    
    p = paragraph_with_styles(doc)
    for item in p:
        print(item["index"], item["style"], item["text"])